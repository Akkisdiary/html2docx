"""
Make 'span' in tags dict a stack
maybe do the same for all tags in case of unclosed tags?
optionally use bs4 to clean up invalid html?

the idea is that there is a method that converts html files into docx
but also have api methods that let user have more control e.g. so they
can nest calls to something like 'convert_chunk' in loops

user can pass existing document object as arg
(if they want to manage rest of document themselves)

How to deal with block level style applied over table elements? e.g. text align
"""
import re
import argparse
import io
import os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx
import docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH

from bs4 import BeautifulSoup

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5  # To stop indents going off the page


def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)


def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])


def fetch_image(url):
    """
    Attempts to fetch an image from a url.
    If successful returns a bytes object, else returns None

    :return:
    """
    try:
        with urllib.request.urlopen(url) as response:
            # security flaw?
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None


def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)


def remove_whitespace(string):
    string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'>\s<', '><', string).strip()


def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def pixels_to_inch(pixels: float):
    return pixels * 0.0104166667


font_styles = {
    'b': {'font-weight': 'bold'},
    'strong': {'font-weight': 'bold'},
    'th': {'font-weight': 'bold'},
    'em': {'font-style': 'italic'},
    'i': {'font-style': 'italic'},
    'u': {'text-decoration': 'underline'},
    's': {'text-decoration': 'line-through'},
    'sup': {'script-style': 'superscript'},
    'sub': {'script-style': 'subscript'},
}

font_tags = ['b', 'strong', 'th', 'em', 'i', 'u', 's', 'sup', 'sub']
block_tags = ['p', 'div']
paragraph_styles = {
    'ol': 'List Number',
    'ul': 'List Bullet',
}


class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.document = Document()
        self.block = None
        self.current_block = None
        self.run_tags = []
        self.runs = []
        self.add_column = True

    def set_initial_attrs(self, document=None):
        if document:
            self.doc = document

    def format_block(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.block.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.block.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.block.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin = int(re.sub(r'[a-z]+', '', margin))
            if units == 'px':
                self.block.paragraph_format.left_indent = Inches(
                    min(margin // 10 * INDENT, MAX_INDENT))
            # TODO handle non px units

    def add_style_to_run(self, style, run):
        if not style:
            return
        if 'color' in style:
            color = re.sub(r'[a-z()]+', '', style['color'])
            colors = [int(x) for x in color.split(',')]
            run.font.color.rgb = RGBColor(
                *colors)
        if 'background-color' in style:
            color = color = re.sub(r'[a-z()]+', '', style['background-color'])
            colors = [int(x) for x in color.split(',')]
            run.font.highlight_color = WD_COLOR.YELLOW
        if 'font-size' in style:
            size = re.sub(r'[a-z()]+', '', style['font-size'])
            run.font.size = Pt(int(size))
        if 'font-weight' in style:
            weight = style['font-weight']
            if weight == 'bold':
                run.bold = True
        if 'font-style' in style:
            sty = style['font-style']
            if sty == 'italic':
                run.italic = True
        if 'text-decoration' in style:
            if style['text-decoration'] == 'underline':
                run.underline = True
            if style['text-decoration'] == 'line-through':
                run.font.strike = True
        if 'script-style' in style:
            if style['script-style'] == 'superscript':
                run.font.superscript = True
            if style['script-style'] == 'subscript':
                run.font.subscript = True

    def parse_dict_string(self, string, separator=';'):
        string = string.replace(' ', '').split(separator)
        parsed_dict = dict([x.split(':') for x in string if ':' in x])
        return parsed_dict

    def add_img(self, current_attrs):
        current_attrs = dict(current_attrs)
        src = current_attrs.get('src')
        if is_url(src):
            try:
                image = fetch_image(src)
            except urllib.error.URLError:
                image = None
        else:
            image = '../writingartist/writing_artist/'+src
        # add image to doc
        if image:
            try:
                if current_attrs.get('style'):
                    style = self.parse_dict_string(current_attrs.get('style'))
                    width = style.get('width')
                    if not width:
                        width = 200
                    if width.endswith('px'):
                        width = float(width.rstrip('px'))
                    width = pixels_to_inch(width)
                # img_run = self.document.add_paragraph().add_run()
                # img_run.add_picture(image, None, None)
                self.document.add_picture(image, width=Inches(width))
            except FileNotFoundError:
                image = None
        if not image:
            if is_url(src):
                self.document.add_paragraph("<image: %s>" % src)
            else:
                # avoid exposing filepaths in document
                self.document.add_paragraph("<image: %s>" %
                                            get_filename_from_url(src))
        # add styles?

    def handle_starttag(self, tag, attrs):
        style = dict(attrs).get('style')
        if style:
            style = self.parse_dict_string(style)
        if tag in block_tags:
            self.current_block = tag
            self.block = self.document.add_paragraph()
            if style:
                self.format_block(style)
            return
        if tag in font_tags:
            style = font_styles[tag]
        if tag == 'table':
            r = self.tables[0]['rows']
            c = self.tables[0]['cols']
            self.block = self.document.add_table(r, c)
            self.current_block = tag
        if tag == 'tr':
            self.tables[0]['curr_row'] += 1
        if tag == 'td':
            curr_table = self.tables[0]
            curr_table['curr_col'] += 1
            if curr_table['curr_col'] >= curr_table['cols']:
                self.tables[0]['curr_col'] = 0
        if tag in ('ol', 'ul'):
            self.current_block = tag
            return
        if tag == 'li':
            self.paragraph_style = paragraph_styles.get(
                self.current_block,
                None,
            )
            self.block = self.document.add_paragraph(
                style=self.paragraph_style)
            if style:
                self.format_block(style)
            return
        self.run_tags.append({'style': style, 'runs': []})

    def handle_data(self, data):
        if self.current_block == 'table':
            r = self.tables[0]['curr_row']
            c = self.tables[0]['curr_col']
            curr_cell = self.block.cell(r, c)
            curr_cell.text = data
            return
        if not self.block:
            self.block = self.document.add_paragraph()
        run = self.block.add_run(data)
        if len(self.run_tags) == 0:
            return
        for i in range(len(self.run_tags)):
            self.run_tags[i]['runs'].append(run)

    def handle_endtag(self, tag):
        if tag in (*block_tags, 'tbody', 'ol', 'ul', 'li'):
            return
        if tag == 'table':
            self.tables.pop(0)
            return
        for run in self.run_tags[-1]['runs']:
            self.add_style_to_run(self.run_tags[-1]['style'], run)
        self.run_tags.pop(-1)

    def handle_startendtag(self, tag, attrs):
        if tag == 'br':
            self.handle_data('\n')
            return
        if tag == 'img':
            self.add_img(attrs)

    def run_process(self, html):
        self.soup = BeautifulSoup(html, 'html.parser')
        self.tables = []
        for table in self.soup.find_all('table', recursive=False):
            rows = len(table.find_all('tr'))
            cols = len(table.find_all('td'))
            self.tables.append({
                'rows': rows,
                'cols': cols//rows,
                'curr_row': -1,
                'curr_col': -1,
            })
        self.feed(remove_whitespace(html))

    def from_file(self, file_path, doc_name=None):
        with open(file_path, 'r') as infile:
            html = infile.read()
        self.run_process(html)
        # if not doc_name:
        #     path, file_name = os.path.split(file_path)
        #     file_name = file_name.split('.')[0]
        #     doc_name = '%s/new_%s' % ('.', file_name)
        self.document.save('./tests/outputs/xxxxx.docx')

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' %
                             docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' %
                             docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated

        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables


if __name__ == '__main__':

    arg_parser = argparse.ArgumentParser(
        description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument(
        'filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx',
        nargs='?',
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]',
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true',
                            help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)
